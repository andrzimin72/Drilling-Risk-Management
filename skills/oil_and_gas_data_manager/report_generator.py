"""
Report Generator
Generates branded Word documents (.docx) with Gazprom Neft styling.
Includes risk heat map, executive summary, and detailed risk register.
Requires: pip install python-docx
"""
from __future__ import annotations
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. Report generation disabled. Install with: pip install python-docx")


# ---------------------------------------------------------------------------
# Brand Colors (Gazprom Neft-inspired)
# ---------------------------------------------------------------------------
class BrandColors:
    PRIMARY_BLUE = RGBColor(0x00, 0x72, 0xCE)      # Gazprom blue
    DARK_BLUE = RGBColor(0x00, 0x4C, 0x97)
    LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
    DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    CRITICAL_RED = RGBColor(0xC0, 0x00, 0x00)
    HIGH_ORANGE = RGBColor(0xFF, 0x6B, 0x00)
    MEDIUM_YELLOW = RGBColor(0xFF, 0xD7, 0x00)
    LOW_GREEN = RGBColor(0x4C, 0xAF, 0x50)


# ---------------------------------------------------------------------------
# Report Generator
# ---------------------------------------------------------------------------
class ReportGenerator:
    """Generates branded Word documents with risk analysis."""

    def __init__(
        self,
        language: str = "ru",
        logo_path: Optional[Path] = None,
        company_name: str = "ПАО «Газпром нефть»",
    ):
        """
        Args:
            language: "ru" for Russian, "en" for English
            logo_path: Optional path to company logo image
            company_name: Company name for cover page
        """
        if not HAS_DOCX:
            raise ImportError("python-docx required. Install with: pip install python-docx")
        self.language = language
        self.logo_path = logo_path
        self.company_name = company_name

    def generate(
        self,
        swarm_result: Any,
        risk_registry: Any,
        output_path: Path,
    ) -> Path:
        """
        Generate complete risk analysis report.
        
        Args:
            swarm_result: SwarmResult from orchestrator
            risk_registry: RiskRegistry with all identified risks
            output_path: Path to save .docx file
            
        Returns:
            Path to generated document
        """
        doc = Document()
        self._setup_document_styles(doc)

        # Build report sections
        self._add_cover_page(doc, swarm_result, risk_registry)
        self._add_executive_summary(doc, risk_registry)
        self._add_risk_matrix(doc, risk_registry)
        self._add_critical_risks(doc, risk_registry)
        self._add_high_risks(doc, risk_registry)
        self._add_medium_low_risks_table(doc, risk_registry)
        self._add_domain_summaries(doc, swarm_result)
        self._add_recommendations(doc, risk_registry)
        self._add_appendix(doc, swarm_result, risk_registry)

        # Save
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"Risk report generated: {output_path}")
        return output_path

    # -----------------------------------------------------------------------
    # Document setup
    # -----------------------------------------------------------------------
    def _setup_document_styles(self, doc: Document) -> None:
        """Configure document styles and margins."""
        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = BrandColors.DARK_GRAY

    # -----------------------------------------------------------------------
    # Cover Page
    # -----------------------------------------------------------------------
    def _add_cover_page(self, doc: Document, swarm_result: Any, registry: Any) -> None:
        """Add branded cover page."""
        # Spacer
        for _ in range(4):
            doc.add_paragraph()

        # Company name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.company_name)
        run.font.size = Pt(16)
        run.font.color.rgb = BrandColors.PRIMARY_BLUE
        run.font.bold = True

        doc.add_paragraph()

        # Logo placeholder (if provided)
        if self.logo_path and self.logo_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(self.logo_path), width=Inches(2.5))
            doc.add_paragraph()

        # Title
        title = (
            "ОТЧЁТ ОБ АНАЛИЗЕ РИСКОВ БУРЕНИЯ"
            if self.language == "ru"
            else "DRILLING RISK ANALYSIS REPORT"
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(24)
        run.font.color.rgb = BrandColors.DARK_BLUE
        run.font.bold = True

        doc.add_paragraph()

        # Well info
        well_name = getattr(swarm_result, "well_name", None) or registry.well_name or "Unknown Well"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self._t('Скважина', 'Well')}: {well_name}")
        run.font.size = Pt(14)
        run.font.color.rgb = BrandColors.DARK_GRAY

        if registry.pad_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{self._t('Куст', 'Pad')}: {registry.pad_name}")
            run.font.size = Pt(14)

        # Spacer
        for _ in range(6):
            doc.add_paragraph()

        # Date and metadata
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self._t('Дата формирования', 'Generated')}: {datetime.utcnow().strftime('%d.%m.%Y')}")
        run.font.size = Pt(11)
        run.font.color.rgb = BrandColors.DARK_GRAY

        summary = registry.get_summary()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"{self._t('Выявлено рисков', 'Total risks identified')}: {summary['total_risks']}"
        )
        run.font.size = Pt(11)

        # Page break
        doc.add_page_break()

    # -----------------------------------------------------------------------
    # Executive Summary
    # -----------------------------------------------------------------------
    def _add_executive_summary(self, doc: Document, registry: Any) -> None:
        """Add executive summary section."""
        self._add_heading(doc, self._t("Краткое резюме", "Executive Summary"), level=1)

        summary = registry.get_summary()
        total = summary["total_risks"]

        if total == 0:
            p = doc.add_paragraph(
                self._t(
                    "В ходе анализа не выявлено существенных рисков. Все показатели в пределах нормы.",
                    "No significant risks identified during analysis. All metrics within acceptable limits.",
                )
            )
            return

        # Summary paragraph
        critical = summary["critical_count"]
        high = summary["high_count"]
        score = summary["risk_score"]

        intro = (
            f"В ходе анализа данных скважины выявлено {total} рисков, "
            f"из них {critical} критических и {high} высоких. "
            f"Средневзвешенная оценка риска: {score}/25."
            if self.language == "ru"
            else f"Analysis identified {total} risks, including {critical} critical and {high} high. "
                 f"Weighted risk score: {score}/25."
        )
        doc.add_paragraph(intro)

        # Risk level breakdown table
        doc.add_paragraph()
        table = doc.add_table(rows=5, cols=3)
        table.style = "Light Grid Accent 1"

        headers = [
            self._t("Уровень риска", "Risk Level"),
            self._t("Количество", "Count"),
            self._t("Доля", "Share"),
        ]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        from .risk_manager import RiskLevel
        levels = [
            (RiskLevel.CRITICAL, BrandColors.CRITICAL_RED),
            (RiskLevel.HIGH, BrandColors.HIGH_ORANGE),
            (RiskLevel.MEDIUM, BrandColors.MEDIUM_YELLOW),
            (RiskLevel.LOW, BrandColors.LOW_GREEN),
        ]
        for row_idx, (level, color) in enumerate(levels, start=1):
            count = summary["by_level"].get(level.value, 0)
            share = f"{count/total*100:.0f}%" if total > 0 else "0%"
            table.rows[row_idx].cells[0].text = level.display_name_ru if self.language == "ru" else level.display_name_en
            table.rows[row_idx].cells[1].text = str(count)
            table.rows[row_idx].cells[2].text = share
            # Color-code the level cell
            for paragraph in table.rows[row_idx].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color
                    run.font.bold = True

        doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Risk Matrix (5x5 Heat Map)
    # -----------------------------------------------------------------------
    def _add_risk_matrix(self, doc: Document, registry: Any) -> None:
        """Add 5x5 risk heat map."""
        self._add_heading(doc, self._t("Матрица рисков", "Risk Matrix"), level=1)

        doc.add_paragraph(
            self._t(
                "Матрица 5×5 отображает распределение рисков по вероятности и влиянию.",
                "The 5×5 matrix shows risk distribution by probability and impact.",
            )
        )

        matrix = registry.get_risk_matrix()

        # Build 6x6 table (header + 5 rows)
        table = doc.add_table(rows=6, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        table.rows[0].cells[0].text = self._t("Вероятность ↓ / Влияние →", "Probability ↓ / Impact →")
        for i in range(1, 6):
            cell = table.rows[0].cells[i]
            cell.text = str(i)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True

        # Probability labels (5 down to 1)
        for row_idx in range(1, 6):
            prob = 6 - row_idx  # 5, 4, 3, 2, 1
            table.rows[row_idx].cells[0].text = str(prob)
            for p in table.rows[row_idx].cells[0].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True

        # Fill matrix cells with counts and colors
        from .risk_manager import RiskLevel
        for row_idx in range(1, 6):
            prob = 6 - row_idx
            for impact in range(1, 6):
                count = matrix.get(prob, {}).get(impact, 0)
                cell = table.rows[row_idx].cells[impact]
                cell.text = str(count) if count > 0 else ""
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(14)
                        run.font.bold = True

                # Color based on risk level
                level = RiskLevel.from_score(prob, impact)
                self._set_cell_background(cell, level.color_hex)

        doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Critical & High Risks (Detailed)
    # -----------------------------------------------------------------------
    def _add_critical_risks(self, doc: Document, registry: Any) -> None:
        """Add critical risks section."""
        critical = registry.get_risks(level=__import__("risk_manager", fromlist=["RiskLevel"]).RiskLevel.CRITICAL)
        if not critical:
            return

        self._add_heading(doc, self._t("Критические риски", "Critical Risks"), level=1)

        intro = (
            f"Выявлено {len(critical)} критических рисков, требующих немедленного внимания."
            if self.language == "ru"
            else f"{len(critical)} critical risks identified requiring immediate attention."
        )
        p = doc.add_paragraph(intro)
        for run in p.runs:
            run.font.color.rgb = BrandColors.CRITICAL_RED
            run.font.bold = True

        for risk in critical:
            self._add_risk_detail(doc, risk)

    def _add_high_risks(self, doc: Document, registry: Any) -> None:
        """Add high risks section."""
        high = registry.get_risks(level=__import__("risk_manager", fromlist=["RiskLevel"]).RiskLevel.HIGH)
        if not high:
            return

        self._add_heading(doc, self._t("Высокие риски", "High Risks"), level=1)

        intro = (
            f"Выявлено {len(high)} высоких рисков, требующих планирования мер снижения."
            if self.language == "ru"
            else f"{len(high)} high risks identified requiring mitigation planning."
        )
        p = doc.add_paragraph(intro)
        for run in p.runs:
            run.font.color.rgb = BrandColors.HIGH_ORANGE
            run.font.bold = True

        for risk in high:
            self._add_risk_detail(doc, risk)

    def _add_risk_detail(self, doc: Document, risk: Any) -> None:
        """Add detailed risk block."""
        # Title with color indicator
        title = risk.title_ru if self.language == "ru" else risk.title_en
        p = doc.add_paragraph()
        run = p.add_run(f"■ {risk.risk_id}: {title}")
        run.font.bold = True
        run.font.size = Pt(12)
        color_map = {
            "critical": BrandColors.CRITICAL_RED,
            "high": BrandColors.HIGH_ORANGE,
            "medium": BrandColors.MEDIUM_YELLOW,
            "low": BrandColors.LOW_GREEN,
        }
        run.font.color.rgb = color_map.get(risk.risk_level.value, BrandColors.DARK_GRAY)

        # Description
        desc = risk.description_ru if self.language == "ru" else risk.description_en
        doc.add_paragraph(desc)

        # Metadata
        meta_text = (
            f"{self._t('Вероятность', 'Probability')}: {risk.probability}/5 | "
            f"{self._t('Влияние', 'Impact')}: {risk.impact}/5 | "
            f"{self._t('Источник', 'Source')}: {risk.source_agent}"
        )
        p = doc.add_paragraph(meta_text)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.italic = True

        # Mitigations
        mitigations = risk.mitigation_ru if self.language == "ru" else risk.mitigation_en
        if mitigations:
            p = doc.add_paragraph(self._t("Рекомендуемые меры:", "Recommended mitigations:"))
            for run in p.runs:
                run.font.bold = True
            for m in mitigations:
                doc.add_paragraph(m, style="List Bullet")

        doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Medium & Low Risks (Table format)
    # -----------------------------------------------------------------------
    def _add_medium_low_risks_table(self, doc: Document, registry: Any) -> None:
        """Add medium and low risks in compact table."""
        from .risk_manager import RiskLevel
        medium = registry.get_risks(level=RiskLevel.MEDIUM)
        low = registry.get_risks(level=RiskLevel.LOW)
        other = medium + low

        if not other:
            return

        self._add_heading(doc, self._t("Средние и низкие риски", "Medium & Low Risks"), level=1)

        table = doc.add_table(rows=len(other) + 1, cols=5)
        table.style = "Light Grid Accent 1"

        headers = [
            "ID",
            self._t("Название", "Title"),
            self._t("Категория", "Category"),
            self._t("Уровень", "Level"),
            self._t("Оценка", "Score"),
        ]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True

        for row_idx, risk in enumerate(other, start=1):
            title = risk.title_ru if self.language == "ru" else risk.title_en
            level_display = risk.risk_level.display_name_ru if self.language == "ru" else risk.risk_level.display_name_en
            score = risk.probability * risk.impact

            table.rows[row_idx].cells[0].text = risk.risk_id
            table.rows[row_idx].cells[1].text = title[:80]
            table.rows[row_idx].cells[2].text = risk.category.value
            table.rows[row_idx].cells[3].text = level_display
            table.rows[row_idx].cells[4].text = f"{score}/25"

    # -----------------------------------------------------------------------
    # Domain Summaries
    # -----------------------------------------------------------------------
    def _add_domain_summaries(self, doc: Document, swarm_result: Any) -> None:
        """Add per-domain summaries from agent results."""
        self._add_heading(doc, self._t("Результаты по областям", "Domain Summaries"), level=1)

        agent_results = getattr(swarm_result, "agent_results", {})
        domain_order = ["drilling", "logs", "completions", "production", "directional", "hse"]

        for domain in domain_order:
            result = agent_results.get(domain)
            if not result:
                continue
            if isinstance(result, dict):
                status = result.get("status", "")
                summary = result.get("summary", "")
            else:
                status = getattr(result, "status", "")
                summary = getattr(result, "summary", "")

            if status in ("skipped", "error") or not summary:
                continue

            self._add_heading(doc, domain.upper(), level=2)
            doc.add_paragraph(summary)

    # -----------------------------------------------------------------------
    # Recommendations
    # -----------------------------------------------------------------------
    def _add_recommendations(self, doc: Document, registry: Any) -> None:
        """Add actionable recommendations section."""
        self._add_heading(doc, self._t("Рекомендации", "Recommendations"), level=1)

        critical = registry.get_risks(level=__import__("risk_manager", fromlist=["RiskLevel"]).RiskLevel.CRITICAL)
        high = registry.get_risks(level=__import__("risk_manager", fromlist=["RiskLevel"]).RiskLevel.HIGH)

        if not critical and not high:
            doc.add_paragraph(
                self._t(
                    "Критических и высоких рисков не выявлено. Продолжать мониторинг в штатном режиме.",
                    "No critical or high risks identified. Continue standard monitoring.",
                )
            )
            return

        doc.add_paragraph(
            self._t(
                "На основе анализа выявленных рисков рекомендуются следующие первоочередные действия:",
                "Based on identified risks, the following priority actions are recommended:",
            )
        )

        # Aggregate top mitigations
        seen_mitigations = set()
        recommendations = []
        for risk in critical + high:
            mitigations = risk.mitigation_ru if self.language == "ru" else risk.mitigation_en
            for m in mitigations[:2]:  # Top 2 per risk
                if m not in seen_mitigations:
                    seen_mitigations.add(m)
                    recommendations.append((risk.risk_id, m))

        for risk_id, rec in recommendations[:15]:  # Cap at 15
            doc.add_paragraph(f"[{risk_id}] {rec}", style="List Number")

    # -----------------------------------------------------------------------
    # Appendix
    # -----------------------------------------------------------------------
    def _add_appendix(self, doc: Document, swarm_result: Any, registry: Any) -> None:
        """Add appendix with metadata and raw flags."""
        doc.add_page_break()
        self._add_heading(doc, self._t("Приложение: Метаданные анализа", "Appendix: Analysis Metadata"), level=1)

        # Swarm metadata
        doc.add_paragraph(self._t("Параметры анализа:", "Analysis parameters:"))
        meta_items = [
            (self._t("ID задачи", "Task ID"), getattr(swarm_result, "task_id", "N/A")),
            (self._t("Обработано файлов", "Files processed"), getattr(swarm_result, "files_processed", 0)),
            (self._t("Успешных агентов", "Agents succeeded"), getattr(swarm_result, "agents_succeeded", 0)),
            (self._t("Время анализа", "Analysis time"), f"{getattr(swarm_result, 'elapsed_seconds', 0):.2f}s"),
            (self._t("Дата отчёта", "Report date"), datetime.utcnow().strftime("%d.%m.%Y %H:%M UTC")),
        ]
        for label, value in meta_items:
            doc.add_paragraph(f"• {label}: {value}")

        # Quality flags summary
        all_flags = getattr(swarm_result, "quality_flags", [])
        if all_flags:
            doc.add_paragraph()
            self._add_heading(doc, self._t("Флаги качества", "Quality Flags"), level=2)
            doc.add_paragraph(f"{self._t('Всего флагов', 'Total flags')}: {len(all_flags)}")
            for flag in all_flags[:20]:
                doc.add_paragraph(flag, style="List Bullet")
            if len(all_flags) > 20:
                doc.add_paragraph(f"... {self._t('и ещё', 'and')} {len(all_flags) - 20}")

    # -----------------------------------------------------------------------
    # Helpers
    # -----------------------------------------------------------------------
    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        """Add styled heading."""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = BrandColors.DARK_BLUE

    def _set_cell_background(self, cell: Any, hex_color: str) -> None:
        """Set cell background color."""
        try:
            hex_color = hex_color.lstrip("#")
            shading_elm = cell._element.get_or_add_tcPr()
            shading = shading_elm.makeelement(qn("w:shd"), {
                qn("w:fill"): hex_color,
                qn("w:val"): "clear",
            })
            shading_elm.append(shading)
        except Exception as e:
            logger.debug(f"Could not set cell background: {e}")

    def _t(self, ru: str, en: str) -> str:
        """Translate helper."""
        return ru if self.language == "ru" else en