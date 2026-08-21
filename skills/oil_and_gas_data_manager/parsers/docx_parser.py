"""
DOCX Document Parser
Parses Word documents (.docx) for oil and gas engineering content.
"""
from __future__ import annotations
from pathlib import Path
from typing import Any

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .pdf_reports import _extract_entities, _bucket_entities, _extract_npt_events

class DocxParser:
    """Parse Word documents for oil and gas engineering data."""
    def extract_text(self, path: Path) -> str:
        if not HAS_DOCX:
            return f"[python-docx not installed — cannot parse {path.name}]"
        try:
            doc = Document(str(path))
            parts = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_vals: parts.append(" | ".join(row_vals))
            return "\n".join(parts)
        except Exception as exc:
            return f"[DOCX READ ERROR: {exc}]"

    def extract_structured(
        self, path: Path, raw_text: str = ""
    ) -> tuple[dict[str, Any], list[dict], list[dict]]:
        if not HAS_DOCX:
            return ({"drilling": {}, "directional": {}, "completions": {}, "logs": {}, "production": {}}, [],
                    [{"source_section": "python_docx_unavailable", "page_or_depth_range": "", "confidence": 0.0}])
                    
        text = raw_text or self.extract_text(path)
        entities = _extract_entities(text)
        extracted = _bucket_entities(entities)
        npt_events = _extract_npt_events(text)
        if npt_events: extracted["drilling"]["npt_events"] = npt_events
        tables = self._extract_tables(path)
        references = [{"source_section": "docx_paragraphs_and_tables", "page_or_depth_range": path.name, "confidence": 0.80}]
        return extracted, tables, references

    def _extract_tables(self, path: Path) -> list[dict[str, Any]]:
        if not HAS_DOCX: return []
        tables = []
        try:
            doc = Document(str(path))
            for tbl_idx, table in enumerate(doc.tables):
                if not table.rows: continue
                header_row = [cell.text.strip() for cell in table.rows[0].cells]
                data_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:] if any(cell.text.strip() for cell in row.cells)]
                if header_row:
                    tables.append({"name": f"table_{tbl_idx + 1}", "columns": header_row, "rows": data_rows[:500]})
        except Exception:
            pass
        return tables
