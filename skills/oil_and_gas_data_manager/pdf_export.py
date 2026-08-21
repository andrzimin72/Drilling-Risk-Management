"""
PDF Export Module
Exports risk reports to PDF with a fallback chain:
  1. docx2pdf (Windows/Mac with Microsoft Word) — highest fidelity
  2. weasyprint (Linux/Mac, headless) — good fidelity
  3. Direct HTML-to-PDF via built-in template — universal fallback
"""
from __future__ import annotations
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Check available backends
try:
    from docx2pdf import convert as docx2pdf_convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

try:
    from weasyprint import HTML as WeasyHTML
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False


def export_to_pdf(
    docx_path: str | Path,
    output_path: str | Path,
    html_fallback: Optional[str] = None,
) -> tuple[bool, str]:
    """
    Convert a Word document to PDF.

    Args:
        docx_path: Path to the source .docx file
        output_path: Path for the output .pdf file
        html_fallback: Optional HTML string for weasyprint fallback

    Returns:
        (success: bool, message: str)
    """
    docx_path = Path(docx_path)
    output_path = Path(output_path)

    if not docx_path.exists():
        return False, f"Source file not found: {docx_path}"

    # Strategy 1: docx2pdf (requires MS Word)
    if HAS_DOCX2PDF:
        try:
            logger.info("Attempting PDF export via docx2pdf (MS Word)...")
            docx2pdf_convert(str(docx_path), str(output_path))
            if output_path.exists():
                logger.info(f"PDF exported via docx2pdf: {output_path}")
                return True, "Exported via Microsoft Word (docx2pdf)"
        except Exception as exc:
            logger.warning(f"docx2pdf failed: {exc}")

    # Strategy 2: weasyprint (headless, Linux-friendly)
    if HAS_WEASYPRINT and html_fallback:
        try:
            logger.info("Attempting PDF export via weasyprint...")
            WeasyHTML(string=html_fallback).write_pdf(str(output_path))
            if output_path.exists():
                logger.info(f"PDF exported via weasyprint: {output_path}")
                return True, "Exported via WeasyPrint (HTML template)"
        except Exception as exc:
            logger.warning(f"weasyprint failed: {exc}")

    # Strategy 3: Generate minimal PDF with pure Python
    try:
        logger.info("Attempting minimal PDF export via pure Python...")
        _generate_minimal_pdf(docx_path, output_path)
        if output_path.exists():
            return True, "Exported via built-in PDF generator (basic formatting)"
    except Exception as exc:
        logger.error(f"All PDF export strategies failed: {exc}")

    return False, (
        "PDF export unavailable. Install one of:\n"
        "  • pip install docx2pdf (requires Microsoft Word)\n"
        "  • pip install weasyprint (requires system libraries: "
        "libpango, libcairo on Linux)"
    )


def _generate_minimal_pdf(docx_path: Path, output_path: Path) -> None:
    """
    Pure-Python PDF generation fallback.
    Extracts text from DOCX and renders a simple PDF.
    Uses fpdf2 if available, otherwise raises ImportError.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError(
            "No PDF backend available. Install: pip install fpdf2"
        )

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required for PDF fallback")

    # Read DOCX content
    doc = Document(str(docx_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                paragraphs.append(row_text)

    # Generate PDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Try to add Unicode font (important for Russian text!)
    try:
        # Common Cyrillic font paths
        font_candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial.ttf",
        ]
        font_path = None
        for fp in font_candidates:
            if Path(fp).exists():
                font_path = fp
                break

        if font_path:
            pdf.add_font("DejaVu", "", font_path, uni=True)
            pdf.set_font("DejaVu", size=10)
        else:
            pdf.set_font("Helvetica", size=10)
    except Exception:
        pdf.set_font("Helvetica", size=10)

    # Write content
    for i, text in enumerate(paragraphs):
        # Bold headings (heuristic: short lines that are likely titles)
        if len(text) < 60 and i < 5:
            try:
                pdf.set_font(size=13)
            except Exception:
                pass
        pdf.multi_cell(0, 6, text)
        pdf.ln(1)
        try:
            pdf.set_font(size=10)
        except Exception:
            pass

    pdf.output(str(output_path))


def generate_html_report(
    risk_summary: dict,
    risks: list[dict],
    well_name: str = "",
    pad_name: str = "",
    language: str = "ru",
) -> str:
    """
    Generate a styled HTML report for weasyprint PDF conversion.
    This serves as the fallback when docx2pdf is unavailable.
    """
    critical = risk_summary.get("critical_count", 0)
    high = risk_summary.get("high_count", 0)
    total = risk_summary.get("total_risks", 0)
    score = risk_summary.get("risk_score", 0)

    company = "ПАО «Газпром нефть»" if language == "ru" else "Gazprom Neft PJSC"
    title = "ОТЧЁТ ОБ АНАЛИЗЕ РИСКОВ БУРЕНИЯ" if language == "ru" else "DRILLING RISK ANALYSIS REPORT"
    well_label = "Скважина" if language == "ru" else "Well"
    pad_label = "Куст" if language == "ru" else "Pad"
    level_labels = {
        "critical": ("Критический", "#C00000"),
        "high": ("Высокий", "#FF6B00"),
        "medium": ("Средний", "#B8860B"),
        "low": ("Низкий", "#4CAF50"),
    }

    risk_rows = ""
    for r in risks[:30]:  # Cap at 30 risks for PDF
        level = r.get("risk_level", "low")
        level_name, level_color = level_labels.get(level, ("Unknown", "#999"))
        title_text = r.get("title_ru", r.get("title_en", ""))
        desc_text = r.get("description_ru", r.get("description_en", ""))[:200]
        risk_rows += f"""
        <tr>
            <td style="border:1px solid #ddd;padding:6px;font-size:9px;">{r.get('risk_id','')}</td>
            <td style="border:1px solid #ddd;padding:6px;font-size:9px;">
                <strong style="color:{level_color}">{level_name}</strong><br/>
                <strong>{title_text}</strong><br/>
                <span style="font-size:8px;color:#555;">{desc_text}</span>
            </td>
            <td style="border:1px solid #ddd;padding:6px;font-size:9px;text-align:center;">
                {r.get('probability',0)}×{r.get('impact',0)}
            </td>
            <td style="border:1px solid #ddd;padding:6px;font-size:9px;">{r.get('source_agent','')}</td>
        </tr>"""

    return f"""
    <!DOCTYPE html>
    <html lang="{language}">
    <head>
        <meta charset="utf-8"/>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
                @top-center {{
                    content: "{company}";
                    font-size: 8pt;
                    color: #004C97;
                }}
                @bottom-center {{
                    content: counter(page) " / " counter(pages);
                    font-size: 8pt;
                    color: #999;
                }}
            }}
            body {{
                font-family: 'DejaVu Sans', Arial, sans-serif;
                font-size: 10pt;
                color: #333;
                line-height: 1.5;
            }}
            .header {{
                text-align: center;
                border-bottom: 3px solid #0072CE;
                padding-bottom: 15px;
                margin-bottom: 25px;
            }}
            .header h1 {{
                color: #004C97;
                font-size: 18pt;
                margin: 5px 0;
            }}
            .header .company {{
                color: #0072CE;
                font-size: 12pt;
                font-weight: bold;
            }}
            .summary-box {{
                background: #f5f7fa;
                border-left: 4px solid #0072CE;
                padding: 12px 16px;
                margin: 15px 0;
                border-radius: 4px;
            }}
            .metric-grid {{
                display: flex;
                gap: 15px;
                margin: 15px 0;
            }}
            .metric-card {{
                flex: 1;
                background: #f9fafb;
                border: 1px solid #e0e0e0;
                border-radius: 6px;
                padding: 10px;
                text-align: center;
            }}
            .metric-card .value {{
                font-size: 22pt;
                font-weight: bold;
                color: #004C97;
            }}
            .metric-card .label {{
                font-size: 8pt;
                color: #666;
                margin-top: 4px;
            }}
            .critical .value {{ color: #C00000; }}
            .high .value {{ color: #FF6B00; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 15px 0;
            }}
            th {{
                background: #004C97;
                color: white;
                padding: 8px;
                font-size: 9pt;
                text-align: left;
            }}
            .footer {{
                margin-top: 30px;
                padding-top: 10px;
                border-top: 1px solid #ddd;
                font-size: 8pt;
                color: #999;
                text-align: center;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <div class="company">{company}</div>
            <h1>{title}</h1>
            <p>
                {well_label}: <strong>{well_name or 'N/A'}</strong>
                {f'&nbsp;&nbsp;|&nbsp;&nbsp;{pad_label}: <strong>{pad_name}</strong>' if pad_name else ''}
            </p>
        </div>

        <div class="metric-grid">
            <div class="metric-card">
                <div class="value">{total}</div>
                <div class="label">{'Всего рисков' if language == 'ru' else 'Total Risks'}</div>
            </div>
            <div class="metric-card critical">
                <div class="value">{critical}</div>
                <div class="label">{'Критических' if language == 'ru' else 'Critical'}</div>
            </div>
            <div class="metric-card high">
                <div class="value">{high}</div>
                <div class="label">{'Высоких' if language == 'ru' else 'High'}</div>
            </div>
            <div class="metric-card">
                <div class="value">{score}</div>
                <div class="label">{'Оценка /25' if language == 'ru' else 'Score /25'}</div>
            </div>
        </div>

        <div class="summary-box">
            <strong>{'Реестр рисков' if language == 'ru' else 'Risk Register'}:</strong>
        </div>

        <table>
            <tr>
                <th>ID</th>
                <th>{'Риск' if language == 'ru' else 'Risk'}</th>
                <th>P×I</th>
                <th>{'Источник' if language == 'ru' else 'Source'}</th>
            </tr>
            {risk_rows}
        </table>

        <div class="footer">
            {'Сформировано системой автоматического анализа рисков' if language == 'ru' else 'Generated by Automated Risk Analysis System'}
            &nbsp;|&nbsp; Multi-Agent Swarm Architecture v1.0
        </div>
    </body>
    </html>
    """