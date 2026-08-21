"""
Tests for the Word report generator.
"""
import pytest
from pathlib import Path
from skills.oil_and_gas_data_manager.risk_manager import (
    Risk, RiskLevel, RiskCategory, RiskRegistry
)


def _make_mock_swarm_result():
    """Create a minimal mock SwarmResult."""
    class MockSwarmResult:
        task_id = "test_task_123"
        well_name = "Test Well 123"
        api = "42-999-12345-0000"
        files_processed = 5
        agents_succeeded = 6
        agents_failed = 0
        elapsed_seconds = 12.5
        agent_results = {
            "drilling": {"status": "success", "summary": "Drilling OK"},
            "logs": {"status": "success", "summary": "Logs OK"},
        }
        unified_report = "Test unified report"
        quality_flags = ["FLAG1", "FLAG2"]
        cross_domain_flags = []
    return MockSwarmResult()


def _make_registry_with_risks():
    """Create a registry with diverse risks."""
    registry = RiskRegistry(well_name="Test Well 123", pad_name="Pad 5")
    registry.add_risk(Risk(
        category=RiskCategory.HSE,
        title_en="Fatality recorded",
        title_ru="Зафиксирован смертельный случай",
        description_en="1 fatality on site",
        description_ru="1 смертельный случай на объекте",
        probability=5, impact=5,
        mitigation_en=["Investigate immediately"],
        mitigation_ru=["Немедленно расследовать"],
    ))
    registry.add_risk(Risk(
        category=RiskCategory.TECHNICAL,
        title_en="High DLS detected",
        title_ru="Высокая интенсивность искривления",
        probability=4, impact=4,
    ))
    registry.add_risk(Risk(
        category=RiskCategory.OPERATIONAL,
        title_en="Moderate NPT",
        probability=3, impact=2,
    ))
    return registry


class TestReportGenerator:
    def test_generate_creates_docx(self, tmp_path):
        try:
            from skills.oil_and_gas_data_manager.report_generator import ReportGenerator
        except ImportError:
            pytest.skip("python-docx not installed")

        generator = ReportGenerator(language="en")
        swarm_result = _make_mock_swarm_result()
        registry = _make_registry_with_risks()

        output_path = tmp_path / "test_report.docx"
        result = generator.generate(swarm_result, registry, output_path)

        assert result.exists()
        assert result.stat().st_size > 1000  # Non-empty file

    def test_generate_russian_language(self, tmp_path):
        try:
            from skills.oil_and_gas_data_manager.report_generator import ReportGenerator
        except ImportError:
            pytest.skip("python-docx not installed")

        generator = ReportGenerator(language="ru")
        swarm_result = _make_mock_swarm_result()
        registry = _make_registry_with_risks()

        output_path = tmp_path / "test_report_ru.docx"
        generator.generate(swarm_result, registry, output_path)

        assert output_path.exists()
        # Verify the file can be opened and contains Russian text
        from docx import Document
        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Газпром" in all_text or "ОТЧЁТ" in all_text or "Скважина" in all_text

    def test_generate_english_language(self, tmp_path):
        try:
            from skills.oil_and_gas_data_manager.report_generator import ReportGenerator
        except ImportError:
            pytest.skip("python-docx not installed")

        generator = ReportGenerator(language="en")
        swarm_result = _make_mock_swarm_result()
        registry = _make_registry_with_risks()

        output_path = tmp_path / "test_report_en.docx"
        generator.generate(swarm_result, registry, output_path)

        from docx import Document
        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DRILLING RISK" in all_text or "Well" in all_text

    def test_empty_registry_generates_report(self, tmp_path):
        try:
            from skills.oil_and_gas_data_manager.report_generator import ReportGenerator
        except ImportError:
            pytest.skip("python-docx not installed")

        generator = ReportGenerator(language="en")
        swarm_result = _make_mock_swarm_result()
        registry = RiskRegistry()  # Empty

        output_path = tmp_path / "empty_report.docx"
        generator.generate(swarm_result, registry, output_path)

        assert output_path.exists()

    def test_report_contains_all_sections(self, tmp_path):
        try:
            from skills.oil_and_gas_data_manager.report_generator import ReportGenerator
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        generator = ReportGenerator(language="en")
        swarm_result = _make_mock_swarm_result()
        registry = _make_registry_with_risks()

        output_path = tmp_path / "sections_test.docx"
        generator.generate(swarm_result, registry, output_path)

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Check for key sections
        assert "Executive Summary" in all_text or "Summary" in all_text
        assert "Critical" in all_text or "CRITICAL" in all_text
        assert "Risk Matrix" in all_text or "Matrix" in all_text

    def test_report_with_logo(self, tmp_path):
        try:
            from skills.oil_and_gas_data_manager.report_generator import ReportGenerator
        except ImportError:
            pytest.skip("python-docx not installed")

        # Create a dummy logo file
        logo_path = tmp_path / "logo.png"
        # Minimal valid PNG (1x1 pixel)
        logo_path.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
            b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00"
            b"\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01"
            b"F\xca\x07\x14\x00\x00\x00\x00IEND\xaeB`\x82"
        )

        generator = ReportGenerator(language="en", logo_path=logo_path)
        swarm_result = _make_mock_swarm_result()
        registry = _make_registry_with_risks()

        output_path = tmp_path / "logo_report.docx"
        generator.generate(swarm_result, registry, output_path)

        assert output_path.exists()